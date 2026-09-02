"""Word 导出：同一 ReportData → docx（python-docx，中文字体设置）。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .generator import ReportData

_EA_FONT = "微软雅黑"
_LATIN_FONT = "Calibri"
_ACCENT = RGBColor(0x0B, 0x5F, 0x8A)


def _set_cn(run, size: float = 10.5, bold: bool | None = None,
            color=None) -> None:
    run.font.name = _LATIN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), _EA_FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _heading(doc, text: str, level: int) -> None:
    p = doc.add_paragraph()
    sizes = {0: 20, 1: 15, 2: 13}
    run = p.add_run(text)
    _set_cn(run, size=sizes.get(level, 12), bold=True, color=_ACCENT)
    p.space_after = Pt(6)


def _para(doc, text: str, *, size: float = 10.5, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_cn(run, size=size)
    run.font.italic = italic


def _bullet(doc, left: str, body: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{left}{body}")
    _set_cn(run)


def to_docx(data: ReportData, path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = _LATIN_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), _EA_FONT)

    # 标题与元信息
    _heading(doc, f"智能交通新闻周报 {data.week}", 0)
    _para(doc, f"{data.week_label} ｜ 生成时间：{data.generated_at} ｜ "
               f"收录 {data.relevant_count} 条 / 来源 {data.source_count} 个",
          size=9, italic=True)

    # 概览
    _heading(doc, "一、本周概览", 1)
    _para(doc, data.overview or "（无）")
    for pt in data.overview_points:
        _bullet(doc, "", pt)
    if data.fallback:
        _para(doc, "（说明：本节为自动降级生成，LLM 服务暂不可用）", size=9)

    # 类别分布
    if data.distribution:
        _heading(doc, "二、类别分布", 1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for cell, text in zip(table.rows[0].cells, ["类别", "篇数", "占比"]):
            cell.text = ""
            _set_cn(cell.paragraphs[0].add_run(text), size=10, bold=True)
        total = sum(d["count"] for d in data.distribution) or 1
        for d in data.distribution:
            row = table.add_row()
            pct = d["count"] / total * 100
            for cell, text in zip(row.cells, [d["name"], str(d["count"]),
                                              f"{pct:.1f}%"]):
                cell.text = ""
                _set_cn(cell.paragraphs[0].add_run(text), size=9.5)

    # 分主题要点
    _heading(doc, "三、分主题要点", 1)
    for t in data.themes:
        _heading(doc, t.get("title", "未命名主题"), 2)
        for it in t.get("items", []):
            title = _item_title(data, it["idx"])
            _bullet(doc, title, f"—— {it.get('note', '')}")

    # TOP5
    _heading(doc, "四、本周 TOP 事件", 1)
    if data.top5:
        for rank, idx in enumerate(data.top5, 1):
            _para(doc, f"{rank}. {_item_title(data, idx)}")
    else:
        _para(doc, "（无）")

    # 厂商与集成商动态
    _heading(doc, "五、厂商与集成商动态", 1)
    vendors = [it for it in data.items
               if any(t.startswith("厂商动态") for t in it["tags"]) or it.get("companies")]
    if vendors:
        for it in sorted(vendors, key=lambda x: x.get("importance") or 0, reverse=True):
            companies = "；".join(it.get("companies") or [])
            _bullet(doc, it["title"],
                    f"{'（' + companies + '）' if companies else ''}")
    else:
        _para(doc, "（本周没有明显的厂商/集成商动态条目）")

    # 趋势与下周关注
    _heading(doc, "六、趋势观察", 1)
    for t in data.trends:
        _bullet(doc, "", t)
    if not data.trends:
        _para(doc, "（无）")
    _heading(doc, "七、下周关注", 1)
    for n in data.next_week:
        _bullet(doc, "", n)
    if not data.next_week:
        _para(doc, "（无）")

    # 附录清单
    _heading(doc, f"附录：本周新闻清单（{len(data.items)} 条）", 1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["#", "标题", "来源", "日期", "标签", "重要度"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = ""
        _set_cn(cell.paragraphs[0].add_run(text), size=10, bold=True)
    for it in data.items:
        row = table.add_row()
        values = [str(it["idx"]), it["title"], it["source_name"],
                  str(it.get("published_at") or "—"),
                  "/".join(it.get("tags") or []),
                  str(it.get("importance") or "—")]
        for cell, text in zip(row.cells, values):
            cell.text = ""
            _set_cn(cell.paragraphs[0].add_run(text), size=9.5)

    _para(doc, "说明：本报告由 newsagent 自动生成，仅限部门内部使用。", size=8, italic=True)

    doc.save(str(path))


def _item_title(data: ReportData, idx: int) -> str:
    for it in data.items:
        if it["idx"] == idx:
            return it["title"]
    return "(未知条目)"
